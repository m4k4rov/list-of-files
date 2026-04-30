import customtkinter as ctk
import os
from datetime import datetime
from CTkTreeview import CTkTreeview
from docx import Document
from docx.shared import Cm
import openpyxl
from openpyxl.styles import Font
import threading
import tkinter.messagebox as messagebox
import shutil
from pathlib import Path

# Настройка внешнего вида и темы
ctk.set_appearance_mode("Dark")  # Варианты: "Light", "Dark"
ctk.set_default_color_theme("blue")  # Варианты: "blue", "green", "dark-blue"

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Настройка окна
        self.title("Смена имени / список")
        self.geometry("1210x720")
        self.resizable(True, True)

        # Создание Tabview (виджет для вкладок)
        self.tabview = ctk.CTkTabview(self, width=1200, height=700)
        self.tabview.pack(pady=0, padx=0)

        # Добавление вкладок
        self.tab_1 = self.tabview.add("Переименование файлов")
        self.tab_2 = self.tabview.add("Список файлов")
        self.tab_3 = self.tabview.add("Организация файлов")

        # Переменные
        self.folder_path = ctk.StringVar()
        self.keyword = ctk.StringVar(value="file")
        self.sort_option = ctk.StringVar(value="Имя файла")

        self.selected_folder = None
        self.files_list = []

        # Настройка содержимого для каждой вкладки
        self._setup_tab_1()
        self._setup_tab_2()
        self._setup_tab_3()

    def _setup_tab_1(self):
        """Содержимое первой вкладки"""
        # Заголовок
        title_label = ctk.CTkLabel(self.tab_1, text="Переименование файлов с порядковым номером", font=("Arial", 16, "bold"))
        title_label.pack(pady=10)

        # Поле для выбора папки
        folder_frame = ctk.CTkFrame(self.tab_1)
        folder_frame.pack(fill="x", padx=20, pady=10)

        ctk.CTkLabel(folder_frame, text="Папка:").pack(side="left", padx=5)
        ctk.CTkEntry(folder_frame, textvariable=self.folder_path, width=300).pack(side="left", padx=5)
        ctk.CTkButton(folder_frame, text="Обзор", command=self.browse_folder).pack(side="right", padx=15)

        # Поле для ключевого слова
        keyword_frame = ctk.CTkFrame(self.tab_1)
        keyword_frame.pack(fill="x", padx=20, pady=10)

        ctk.CTkLabel(keyword_frame, text="Ключевое слово:").pack(side="left", padx=5)
        ctk.CTkEntry(keyword_frame, textvariable=self.keyword, width=200).pack(side="left", padx=5)

        # Выбор критерия сортировки
        sort_frame = ctk.CTkFrame(self.tab_1)
        sort_frame.pack(fill="x", padx=20, pady=10)

        ctk.CTkLabel(sort_frame, text="Сортировка по:").pack(side="left", padx=5)
        ctk.CTkComboBox(sort_frame,
                        values=["Имя файла", "Дата изменения", "Дата создания"],
                        variable=self.sort_option).pack(side="left", padx=5)

        # Кнопка запуска переименования
        rename_button = ctk.CTkButton(self.tab_1, text="Переименовать файлы", command=self.rename_files)
        rename_button.pack(pady=20)

        # Область для вывода лога
        self.log_text = ctk.CTkTextbox(self.tab_1, height=150)
        self.log_text.pack(fill="both", expand=True, padx=20, pady=10)
        self.log_text.configure(text_color="white")

    def browse_folder(self):
        folder = ctk.filedialog.askdirectory()
        if folder:
            self.folder_path.set(folder)
            self.log(f"Выбрана папка: {folder}")

    def log(self, message):
        timestamp = datetime.now().strftime("%H:%M:%S")
        self.log_text.insert("end", f"[{timestamp}] {message}\n")
        self.log_text.see("end")

    def rename_files(self):
        folder = self.folder_path.get()
        keyword = self.keyword.get().strip()
        sort_by = self.sort_option.get()

        if not folder or not os.path.isdir(folder):
            self.log("Ошибка: укажите корректную папку!")
            return

        try:
            # Получаем список файлов (исключая каталоги)
            files = [f for f in os.listdir(folder) if os.path.isfile(os.path.join(folder, f))]

            if not files:
                self.log("В папке нет файлов для переименования.")
                return

            # Сортируем файлы в зависимости от выбранного критерия
            if sort_by == "Имя файла":
                files.sort()
            elif sort_by == "Дата изменения":
                files.sort(key=lambda x: os.path.getmtime(os.path.join(folder, x)))
            elif sort_by == "Дата создания":
                files.sort(key=lambda x: os.path.getctime(os.path.join(folder, x)))

            # Переименовываем файлы
            for idx, filename in enumerate(files, start=1):
                old_path = os.path.join(folder, filename)
                file_ext = os.path.splitext(filename)[1]

                # Формируем новое имя
                if keyword:
                    new_name = f"{keyword}_{idx:03d}{file_ext}"
                else:
                    new_name = f"{idx:03d}{file_ext}"

                new_path = os.path.join(folder, new_name)

                os.rename(old_path, new_path)
                self.log(f"Переименован: {filename} → {new_name}")

            self.log("Переименование завершено!")

        except Exception as e:
            self.log(f"Ошибка: {str(e)}")


    def _setup_tab_2(self):
        """Содержимое второй вкладки"""
        # Заголовок
        title_label = ctk.CTkLabel(
            self.tab_2,
            text="📁 Список файлов",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(0, 20))

        # Метка пути
        self.path_label = ctk.CTkLabel(
            self.tab_2,
            text="Выбранная папка: не выбрана",
            font=ctk.CTkFont(size=14),
            wraplength=1100
        )
        self.path_label.pack(fill="x", pady=10)

        # Кнопки управления
        button_frame = ctk.CTkFrame(self.tab_2, fg_color="transparent")
        button_frame.pack(fill="x", pady=10)

        self.select_button = ctk.CTkButton(
            button_frame,
            text="📂 Выбрать папку",
            command=self.select_folder,
            font=ctk.CTkFont(size=14),
            height=40
        )
        self.select_button.pack(side="left", padx=(0, 10))

        export_frame = ctk.CTkFrame(button_frame, fg_color="transparent")
        export_frame.pack(side="right")

        self.docx_button = ctk.CTkButton(
            export_frame,
            text="💾 Экспорт в DOCX",
            command=self.export_to_docx,
            font=ctk.CTkFont(size=12),
            height=35,
            fg_color="#1e88e5",
            hover_color="#1565c0"
        )
        self.docx_button.pack(side="left", padx=5)

        self.xlsx_button = ctk.CTkButton(
            export_frame,
            text="💾 Экспорт в XLSX",
            command=self.export_to_xlsx,
            font=ctk.CTkFont(size=12),
            height=35,
            fg_color="#43a047",
            hover_color="#2e7d32"
        )
        self.xlsx_button.pack(side="left", padx=5)

        # Индикатор прогресса
        self.progress = ctk.CTkProgressBar(self.tab_2)
        self.progress.pack(fill="x", pady=10, padx=10)
        self.progress.set(0)
        self.progress.pack_forget()  # Скрываем изначально

        # Treeview с скроллбаром
        tree_frame = ctk.CTkFrame(self.tab_2)
        tree_frame.pack(fill="both", expand=True, pady=10)
        columns = ('name', 'type', 'size', 'creation', 'modification', 'access')
        self.tree = CTkTreeview(
            tree_frame,
            columns=columns,
            show='headings',
            height=20
        )
        # Заголовки колонок
        self.tree.heading('name', text='Имя файла/папки')
        self.tree.heading('type', text='Тип')
        self.tree.heading('size', text='Размер')
        self.tree.heading('creation', text='Создан')
        self.tree.heading('modification', text='Изменён')
        self.tree.heading('access', text='Открыт')
        # Ширина колонок
        self.tree.column('name', width=500)
        self.tree.column('type', width=80)
        self.tree.column('size', width=100)
        self.tree.column('creation', width=150)
        self.tree.column('modification', width=150)
        self.tree.column('access', width=150)
        # Скроллбар
        scrollbar = ctk.CTkScrollbar(tree_frame, command=self.tree.yview)
        self.tree.configure(yscrollcommand=scrollbar.set)
        self.tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")
        # Строка статуса
        self.status_label = ctk.CTkLabel(
            self.tab_2,
            text="Выберите папку для отображения файлов",
            font=ctk.CTkFont(size=12)
        )
        self.status_label.pack(fill="x", pady=5)

    def select_folder(self):
        """Открывает диалоговое окно для выбора папки и обновляет список файлов."""
        folder_path = ctk.filedialog.askdirectory(title="Выберите папку")
        if folder_path:
            self.path_label.configure(text=f"Выбранная папка: {folder_path}")
            # Запускаем обновление списка в отдельном потоке
            threading.Thread(target=self.update_file_list, args=(folder_path,), daemon=True).start()
    def get_file_info(self, file_path):
        """Получает информацию о файле: размер, даты создания, изменения, открытия."""
        try:
            stat = os.stat(file_path)
            # Размер файла в байтах, преобразуем в КБ или МБ
            size_bytes = stat.st_size
            if size_bytes < 1024:
                size = f"{size_bytes} Б"
            elif size_bytes < 1024**2:
                size = f"{size_bytes / 1024:.2f} КБ"
            else:
                size = f"{size_bytes / (1024**2):.2f} МБ"
            # Даты
            creation_time = datetime.fromtimestamp(stat.st_ctime)
            modification_time = datetime.fromtimestamp(stat.st_mtime)
            access_time = datetime.fromtimestamp(stat.st_atime)
            return {
                'size': size,
                'creation': creation_time.strftime('%Y-%m-%d %H:%M:%S'),
                'modification': modification_time.strftime('%Y-%m-%d %H:%M:%S'),
                'access': access_time.strftime('%Y-%m-%d %H:%M:%S')
            }
        except Exception as e:
            print(f"Ошибка: {str(e)}")
            return {'size': 'N/A', 'creation': 'N/A', 'modification': 'N/A', 'access': 'N/A'}
    def collect_all_files(self, folder_path, progress_callback=None):
        """Рекурсивно собирает все файлы и папки с их метаданными."""
        all_items = []
        total_items = 0
        # Первый проход: подсчёт общего количества элементов
        for root, dirs, files in os.walk(folder_path):
            total_items += len(dirs) + len(files)
        processed_items = 0
        # Второй проход: сбор информации с обновлением прогресса
        for root, dirs, files in os.walk(folder_path):
            # Обработка папок
            for dir_name in dirs:
                dir_path = os.path.join(root, dir_name)
                relative_path = os.path.relpath(dir_path, folder_path)
                all_items.append({
                    'name': relative_path,
                    'type': 'Папка',
                    'size': '-',
                    'creation': '-',
                    'modification': '-',
                    'access': '-'
                })
                processed_items += 1
                if progress_callback:
                    progress_callback(processed_items, total_items)

            # Обработка файлов
            for file_name in files:
                file_path = os.path.join(root, file_name)
                relative_path = os.path.relpath(file_path, folder_path)
                info = self.get_file_info(file_path)
                all_items.append({
                    'name': relative_path,
                    'type': 'Файл',
                    'size': info['size'],
                    'creation': info['creation'],
                    'modification': info['modification'],
                    'access': info['access']
                })
                processed_items += 1
                if progress_callback:
                    progress_callback(processed_items, total_items)

        return all_items

    def update_file_list(self, folder_path):
        """Обновляет список файлов в Treeview с метаданными (рекурсивно)."""
        def update_progress(current, total):
            """Обновляет индикатор прогресса."""
            if total > 0:
                self.progress.set(current / total)
            self.status_label.configure(
                text=f"Обработано: {current}/{total} элементов"
            )
            self.update_idletasks()

        # Очищаем текущий список
        for item in self.tree.get_children():
            self.tree.delete(item)

        try:
            # Показываем прогрессбар
            self.progress.pack(fill="x", padx=10, pady=10)
            self.progress.set(0)

            all_items = self.collect_all_files(folder_path, update_progress)
            file_count = sum(1 for item in all_items if item['type'] == 'Файл')
            folder_count = sum(1 for item in all_items if item['type'] == 'Папка')


            # Сортируем по имени для удобства
            all_items.sort(key=lambda x: x['name'])

            for item in all_items:
                self.tree.insert('', 'end', values=(
                    item['name'],
                    item['type'],
                    item['size'],
                    item['creation'],
                    item['modification'],
                    item['access']
                ))

            self.status_label.configure(
                text=f"Найдено: {len(all_items)} элементов "
                f"({file_count} файлов, {folder_count} папок)"
            )

            # Скрываем прогрессбар после завершения
            self.progress.pack_forget()

        except PermissionError:
            self.status_label.configure(text="Ошибка: нет доступа к папке")
            self.progress.pack_forget()
        except Exception as e:
            self.status_label.configure(text=f"Ошибка: {str(e)}")
            self.progress.pack_forget()

    def export_to_docx(self):
        """Экспортирует список файлов в формат DOCX."""
        if not self.tree.get_children():
            self.status_label.configure(text="Нет данных для экспорта")
            return

        file_path = ctk.filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word документы", "*.docx"), ("Все файлы", "*.*")]
        )

        if file_path:
            doc = Document()

            section = doc.sections[0]
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

            doc.add_heading('Список файлов', 0)

            # Добавляем путь к папке
            #folder_text = self.path_label.cget("text")
            #doc.add_paragraph(f"Папка: {folder_text.replace('Выбранная папка: ', '')}")
            doc.add_paragraph()

            # Создаём таблицу
            headers = ['Имя', 'Тип', 'Размер', 'Дата создания', 'Дата изменения', 'Дата открытия']
            table = doc.add_table(rows=1, cols=len(headers))

            table.style = 'Table Grid'
            table.autofit = False
            table.allow_autofit = False
            #table.width = Cm(18)
            table.columns[0].width = Cm(4)


            # Заголовки таблицы
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header

            # Данные таблицы
            for row in self.tree.get_children():
                values = self.tree.item(row)['values']
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
            doc.save(file_path)
            self.status_label.configure(text=f"Данные экспортированы в DOCX: {file_path}")

    def export_to_xlsx(self):
        """Экспортирует список файлов в формат XLSX."""
        if not self.tree.get_children():
            self.status_label.configure(text="Нет данных для экспорта")
            return

        file_path = ctk.filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel файлы", "*.xlsx"), ("Все файлы", "*.*")]
        )

        if file_path:
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Список файлов"
            ws.column_dimensions['A'].width = 80
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 20
            ws.column_dimensions['E'].width = 20
            ws.column_dimensions['F'].width = 20

            # Заголовки
            headers = ['Имя', 'Тип', 'Размер', 'Дата создания', 'Дата изменения', 'Дата открытия']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)

            # Данные
            row_idx = 2
            for tree_row in self.tree.get_children():
                values = self.tree.item(tree_row)['values']
                for col_idx, value in enumerate(values, 1):
                    ws.cell(row=row_idx, column=col_idx, value=str(value))
                row_idx += 1

            wb.save(file_path)
            self.status_label.configure(text=f"Данные экспортированы в XLSX: {file_path}")

    def _setup_tab_3(self):
        # Заголовок
        title_label = ctk.CTkLabel(
            self.tab_3,
            text="Organizer файлов по расширениям",
            font=ctk.CTkFont(size=20, weight="bold")
        )
        title_label.pack(pady=20)

        # Кнопка выбора папки
        select_folder_btn = ctk.CTkButton(
            self.tab_3,
            text="Выбрать папку для сканирования",
            command=self.select_folder_org,
            width=200
        )
        select_folder_btn.pack(pady=10)

        # Метка выбранной папки
        self.folder_label = ctk.CTkLabel(self.tab_3, text="Папка не выбрана", text_color="gray")
        self.folder_label.pack(pady=5)

        # Фрейм для списка файлов
        files_frame = ctk.CTkFrame(self.tab_3)
        files_frame.pack(fill="both", expand=True, padx=20, pady=10)

        # Текстовое поле для вывода файлов
        self.files_textbox = ctk.CTkTextbox(files_frame, wrap="word")
        self.files_textbox.pack(fill="both", expand=True, padx=10, pady=10)

        # Строка прогресса
        progress_frame = ctk.CTkFrame(self.tab_3, fg_color="transparent")
        progress_frame.pack(fill="x", padx=20, pady=5)

        self.progress_label = ctk.CTkLabel(progress_frame, text="Готов к работе")
        self.progress_label.pack(side="left")

        self.progress_bar = ctk.CTkProgressBar(progress_frame)
        self.progress_bar.pack(side="right", fill="x", expand=True, padx=(10, 0))
        self.progress_bar.set(0)

        # Кнопки действий
        buttons_frame = ctk.CTkFrame(self.tab_3, fg_color="transparent")
        buttons_frame.pack(fill="x", padx=20, pady=10)

        scan_btn = ctk.CTkButton(
            buttons_frame,
            text="Сканировать папку",
            command=self.scan_folder,
            fg_color="green"
        )
        scan_btn.pack(side="left", padx=5)

        organize_btn = ctk.CTkButton(
            buttons_frame,
            text="Организовать файлы",
            command=self.start_organize_files,
            fg_color="orange"
        )
        organize_btn.pack(side="right", padx=5)

    def select_folder_org(self):
        """Выбор папки через диалоговое окно с автоматическим сканированием"""
        folder = ctk.filedialog.askdirectory()
        if folder:
            self.selected_folder = folder
            self.folder_label.configure(text=f"Выбрана: {folder}")
            self.files_textbox.delete("1.0", "end")
            self.scan_folder()  # Автоматическое сканирование после выбора

    def scan_folder(self):
        """Сканирование выбранной папки и всех подпапок"""
        if not self.selected_folder:
            self.show_error("Сначала выберите папку!")
            return

        try:
            self.files_list = []
            total_files = 0

            # Рекурсивный обход всех папок
            for root, dirs, files in os.walk(self.selected_folder):
                for file in files:
                    file_path = os.path.join(root, file)
                    self.files_list.append(file_path)
                    total_files += 1

            # Вывод результатов
            self.files_textbox.delete("1.0", "end")
            self.files_textbox.insert("1.0", f"Найдено файлов: {total_files}\n\n")

            if self.files_list:
                for file_path in self.files_list[:100]:  # Ограничение для отображения
                    self.files_textbox.insert("end", f"{file_path}\n")

                if len(self.files_list) > 100:
                    self.files_textbox.insert("end", f"\n... и ещё {len(self.files_list) - 100} файлов")
            else:
                self.files_textbox.insert("end", "В папке нет файлов.")

        except Exception as e:
            self.show_error(f"Ошибка при сканировании: {str(e)}")

    def start_organize_files(self):
        """Запуск организации файлов в отдельном потоке"""
        if not self.files_list:
            self.show_error("Сначала выполните сканирование!")
            return

        # Запуск в отдельном потоке, чтобы не блокировать интерфейс
        thread = threading.Thread(target=self.organize_files)
        thread.daemon = True
        thread.start()

    def organize_files(self):
        """Организация файлов по расширениям с прогрессом"""
        self.update_progress("Начинаем организацию файлов...", 0)

        try:
            # Создание целевой папки
            target_folder = os.path.join(self.selected_folder, "Organized_Files")
            os.makedirs(target_folder, exist_ok=True)

            # Словарь для группировки по расширениям
            files_by_extension = {}

            # Группировка файлов
            for file_path in self.files_list:
                file_name = os.path.basename(file_path)
                file_ext = Path(file_name).suffix.lower()

                # Если нет расширения, помещаем в отдельную категорию
                if not file_ext:
                    file_ext = "no_extension"
                else:
                    # Убираем точку из расширения
                    file_ext = file_ext[1:]

                if file_ext not in files_by_extension:
                    files_by_extension[file_ext] = []
                files_by_extension[file_ext].append(file_path)


            # Подсчёт общего количества файлов для прогресса
            total_files = len(self.files_list)
            copied_count = 0

            # Копирование файлов с обновлением прогресса
            for ext, files in files_by_extension.items():
                # Создаём папку для расширения
                ext_folder = os.path.join(target_folder, ext)
                os.makedirs(ext_folder, exist_ok=True)

                for file_path in files:
                    try:
                        # Генерируем уникальное имя при конфликте
                        file_name = os.path.basename(file_path)
                        target_path = os.path.join(ext_folder, file_name)

                        # Если файл уже существует, добавляем суффикс
                        counter = 1
                        while os.path.exists(target_path):
                            name_without_ext = Path(file_name).stem
                            extension = Path(file_name).suffix
                            new_name = f"{name_without_ext}_{counter}{extension}"
                            target_path = os.path.join(ext_folder, new_name)
                            counter += 1

                        shutil.copy2(file_path, target_path)
                        copied_count += 1

                        # Обновление прогресса
                        progress = copied_count / total_files
                        self.update_progress(f"Копируется: {file_name} ({copied_count}/{total_files})", progress)

                    except Exception as e:
                        print(f"Ошибка копирования {file_path}: {e}")

            # Завершение процесса
            self.update_progress(f"Завершено! Скопировано {copied_count} файлов в '{target_folder}'", 1.0)
            self.show_info(f"Успешно скопировано {copied_count} файлов в папку '{target_folder}'")

        except Exception as e:
            self.show_error(f"Ошибка при организации файлов: {str(e)}")
        finally:
            # Сбрасываем прогресс-бар после завершения
            self.after(0, lambda: self.progress_bar.set(0))
            # Можно добавить небольшую задержку перед сбросом, если нужно
            # self.after(2000, lambda: self.progress_bar.set(0))

    def update_progress(self, message, progress_value):
        """Обновление строки прогресса и метки"""
        def update():
            self.progress_label.configure(text=message)
            self.progress_bar.set(progress_value)
            # Обновляем интерфейс
            self.update()

        # Выполняем обновление в основном потоке GUI
        self.after(0, update)

    def show_error(self, message):
        """Показать сообщение об ошибке"""
        messagebox.showerror("Ошибка", message)

    def show_info(self, message):
        """Показать информационное сообщение"""
        messagebox.showinfo("Информация", message)


# Запуск приложения
if __name__ == "__main__":
    app = App()
    app.mainloop()
